#!/usr/bin/env python3
"""将 AI学习笔记.md 转换为 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import os

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """给表格单元格设置背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    """添加格式化的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(header))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')
    return table

def add_code_block(doc, code_text, font_size=Pt(7.5)):
    """添加代码块（灰底等宽字体）"""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # 灰底
        shading = p._element.get_or_add_pPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F0F0F0',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

def add_heading_custom(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_text(doc, text, bold=False, italic=False, font_size=Pt(10.5), color=None, indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = font_size
    if color:
        run.font.color.rgb = color
    return p

def add_blockquote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return p

def add_horizontal_rule(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def is_inside_code_block(lines, idx):
    """检查某行是否在代码块内"""
    fence_count = 0
    for i in range(idx + 1):
        if lines[i].strip().startswith('```'):
            fence_count += 1
    return fence_count % 2 == 1

def parse_markdown_table_line(line):
    """解析markdown表格行"""
    # 去掉首尾的 |
    line = line.strip().strip('|')
    cells = [c.strip() for c in line.split('|')]
    return cells


# ============================================================
# 主转换逻辑
# ============================================================

def convert_md_to_docx(md_path, docx_path):
    """将markdown文件转换为Word文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_headers = []
    table_rows = []
    table_is_separator = re.compile(r'^\|[\s\-:|]+\|$')

    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            if in_table:
                # 结束表格
                if table_headers:
                    add_styled_table(doc, table_headers, table_rows)
                table_headers = []
                table_rows = []
                in_table = False
            if in_code_block:
                code_buffer.append('')
            else:
                doc.add_paragraph('')
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, '\n'.join(code_buffer))
                doc.add_paragraph('')
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level <= 3:
                add_heading_custom(doc, text, level)
            else:
                add_heading_custom(doc, text, level)
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # 引用块 >
        if line.strip().startswith('>'):
            # 收集连续引用
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_blockquote(doc, ' '.join(quote_lines))
            continue

        # 表格
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if in_table:
                # 检查是否是分隔行
                if table_is_separator.match(line.strip()):
                    i += 1
                    continue
                # 数据行
                cells = parse_markdown_table_line(line)
                table_rows.append(cells)
            else:
                in_table = True
                table_headers = parse_markdown_table_line(line)
                table_rows = []
            i += 1
            continue
        else:
            if in_table:
                # 表格结束
                if table_headers:
                    add_styled_table(doc, table_headers, table_rows)
                table_headers = []
                table_rows = []
                in_table = False

        # 无序列表
        ul_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if ul_match:
            indent_level = len(ul_match.group(1))
            text = ul_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            # 清除默认run, 添加新run
            for run in p.runs:
                run.text = ''
            run = p.add_run(strip_bold_markers(text))
            run.font.size = Pt(10)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+[\.\、]\s+(.+)$', line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            for run in p.runs:
                run.text = ''
            run = p.add_run(strip_bold_markers(text))
            run.font.size = Pt(10)
            i += 1
            continue

        # 普通段落 - 处理行内格式
        p = doc.add_paragraph()
        process_inline_markdown(p, line)
        i += 1

    # 处理文件末尾未关闭的表格
    if in_table and table_headers:
        add_styled_table(doc, table_headers, table_rows)

    # 保存
    doc.save(docx_path)
    print(f'Word 文档已保存到: {docx_path}')


def strip_bold_markers(text):
    """去掉 ** 标记，返回纯文本"""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)

def process_inline_markdown(paragraph, text):
    """处理行内markdown格式"""
    # 简单处理: 分段 bold / italic / code / plain
    # pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0
    for m in pattern.finditer(text):
        # 前面的纯文本
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            run = paragraph.add_run(plain)
            run.font.size = Pt(10)
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(10)
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(10)
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last_end = m.end()
    # 剩余纯文本
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(10)


if __name__ == '__main__':
    md_file = r'D:\ClaudeCodeTest\test5\AI学习笔记.md'
    docx_file = r'D:\ClaudeCodeTest\test5\AI学习笔记.docx'
    convert_md_to_docx(md_file, docx_file)
